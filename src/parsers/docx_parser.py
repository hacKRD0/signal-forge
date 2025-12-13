"""DOCX document parser.

This module provides functionality to extract text content from Microsoft Word
DOCX files, including text from paragraphs and tables.
"""

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.opc.exceptions import PackageNotFoundError


logger = logging.getLogger(__name__)


def parse_docx(file_path: str) -> str:
    """Extract text content from a DOCX file.

    Reads a Microsoft Word DOCX file and extracts all text content from
    paragraphs and tables. The text is concatenated with newlines to
    preserve document structure.

    Args:
        file_path: Path to the DOCX file to parse (absolute or relative).

    Returns:
        Extracted text as a single string with newlines separating paragraphs
        and table cells. Returns empty string if the document has no text.

    Raises:
        FileNotFoundError: If the specified file does not exist.
        ValueError: If the file is corrupted or not a valid DOCX file.
        PermissionError: If the file cannot be read due to permissions.

    Example:
        >>> text = parse_docx('documents/report.docx')
        >>> print(f"Extracted {len(text)} characters")
    """
    # Convert to Path for better path handling
    path = Path(file_path)

    # Check if file exists
    if not path.exists():
        logger.error(f"File not found: {file_path}")
        raise FileNotFoundError(f"File not found: {file_path}")

    logger.info(f"Parsing DOCX file: {file_path}")

    try:
        # Open the document
        doc = Document(str(path))

        # Extract text from paragraphs
        text_parts = []

        # Extract paragraph text
        logger.debug(f"Extracting text from {len(doc.paragraphs)} paragraphs")
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text_parts.append(paragraph.text)

        # Extract text from tables
        logger.debug(f"Extracting text from {len(doc.tables)} tables")
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))

        # Join all text parts with newlines
        extracted_text = "\n".join(text_parts)

        logger.info(
            f"Successfully extracted {len(extracted_text)} characters "
            f"from {file_path}"
        )

        return extracted_text

    except PackageNotFoundError as e:
        logger.error(f"Corrupted or invalid DOCX file: {file_path}")
        raise ValueError(
            f"File is corrupted or not a valid DOCX file: {file_path}"
        ) from e
    except PermissionError as e:
        logger.error(f"Permission denied reading file: {file_path}")
        raise
    except Exception as e:
        logger.error(f"Unexpected error parsing DOCX file {file_path}: {e}")
        raise ValueError(
            f"Error parsing DOCX file {file_path}: {str(e)}"
        ) from e
